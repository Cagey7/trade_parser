import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor
from db.database import connect_to_db
from doc_gen import generate_data_for_doc
from db.data.month_ranges import month_ranges

# Подключение к БД и загрузка данных
conn = connect_to_db()
year = 2025
country = "Россия"
region = "Республика Казахстан"
units_of_account = "тыс"
month, data_for_doc = generate_data_for_doc(conn, year, country, region, units_of_account)

# Установка стиля текста
def set_run_style(run):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(14)

# Форматирование абзаца
def format_paragraph(paragraph, first_line_indent=True):
    if first_line_indent:
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
    paragraph.paragraph_format.line_spacing = 1

# Заголовок документа
def add_document_header(doc, header_text):
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_paragraph(paragraph, first_line_indent=False)
    run = paragraph.add_run(header_text)
    run.bold = True
    set_run_style(run)

# Обычный параграф
def add_summary_paragraph(doc, text):
    paragraph = doc.add_paragraph()
    format_paragraph(paragraph)
    run = paragraph.add_run(text)
    set_run_style(run)

# Цвет фона ячейки
def set_cell_background(cell, color_hex):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

# Добавление вертикальных отступов (space above/below в ячейке)
def add_cell_vertical_padding(cell, space_pts=4):
    for paragraph in cell.paragraphs:
        paragraph.paragraph_format.space_before = Pt(space_pts)
        paragraph.paragraph_format.space_after = Pt(space_pts)

# Получение доступной ширины страницы
def get_available_page_width(doc):
    section = doc.sections[0]
    return section.page_width - section.left_margin - section.right_margin

# Генерация таблицы с заголовком
def add_summary_table(doc, title, table_data):
    # Заголовок таблицы без отступа после
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_after = Pt(0)  # убираем отступ после заголовка
    format_paragraph(paragraph, first_line_indent=False)
    run = paragraph.add_run(title)
    run.bold = True
    set_run_style(run)

    rows = len(table_data)
    cols = len(table_data[0])
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    available_width = get_available_page_width(doc)
    col_width = available_width / cols

    for i, row in enumerate(table_data):
        for j, cell_text in enumerate(row):
            cell = table.cell(i, j)
            cell.text = str(cell_text)
            cell.width = col_width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0]
                set_run_style(run)

                if i == 0 and j != 0:
                    run.bold = True
                elif i == 1:
                    run.bold = True

                if j == 3 and '-' in cell_text:
                    run.font.color.rgb = RGBColor(255, 0, 0)
                if j == 3 and "ухудшился" in cell_text:
                    run.italic = True

            add_cell_vertical_padding(cell, space_pts=2)

            if i == 0:
                set_cell_background(cell, "D9D9D9")



def add_import_analysis_text(doc, text_blocks):
    # Отступ перед блоком текста
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    for i, text in enumerate(text_blocks):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        paragraph.paragraph_format.line_spacing = 1
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if i == 0:
            # Найти первое слово
            first_word_match = re.match(r'\S+', text)
            first_word = first_word_match.group(0) if first_word_match else ''
            rest_text = text[len(first_word):]

            # Найти последнее число (включая формат с пробелами, запятыми и точками)
            last_number_match = list(re.finditer(r'\d[\d\s.,]*', text))
            if last_number_match:
                last_number = last_number_match[-1]
                split_index = last_number.start()
                middle_text = text[len(first_word):split_index]
                end_text = text[split_index:]
            else:
                middle_text = rest_text
                end_text = ''

            # Жирное первое слово
            run_first = paragraph.add_run(first_word)
            set_run_style(run_first)
            run_first.bold = True

            # Средняя часть — обычная
            run_middle = paragraph.add_run(middle_text)
            set_run_style(run_middle)

            # Конец — жирный (включая последнюю цифру)
            run_end = paragraph.add_run(end_text)
            set_run_style(run_end)
            run_end.bold = True

        elif i == 3 and ':' in text:
            before_colon, after_colon = text.split(':', 1)

            run_bold = paragraph.add_run(before_colon + ':')
            set_run_style(run_bold)
            run_bold.bold = True

            run_rest = paragraph.add_run(after_colon)
            set_run_style(run_rest)

        else:
            run = paragraph.add_run(text)
            set_run_style(run)
            if i in [1, 2]:
                run.italic = True


def generate_export_doc_with_repeating_header(doc, table_header, table_data, month, year, units_of_account):
    # Удаляем 10-й элемент, если есть
    for row in table_data:
        if len(row) > 9:
            del row[9]

    # --- Вспомогательные функции ---
    def merge_cells_vertically(table, col_idx, row_start, row_end):
        start_cell = table.cell(row_start, col_idx)
        for row in range(row_start + 1, row_end + 1):
            start_cell.merge(table.cell(row, col_idx))

    def merge_cells_horizontally(table, row_idx, col_start, col_end):
        start_cell = table.cell(row_idx, col_start)
        for col in range(col_start + 1, col_end + 1):
            start_cell.merge(table.cell(row_idx, col))

    def set_cell_background(cell, color="D9D9D9"):
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), color)
        tcPr.append(shd)

    def center_cell(cell):
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def make_bold(cell):
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    def set_repeat_table_header(row):
        tr = row._tr
        trPr = tr.get_or_add_trPr()
        tblHeader = OxmlElement('w:tblHeader')
        trPr.append(tblHeader)

    # --- Добавление отступа и заголовка ---
    doc.add_paragraph().paragraph_format.space_after = 0

    header_paragraph = doc.add_paragraph(table_header)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_paragraph.runs[0]
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    header_paragraph.paragraph_format.space_after = 0

    # --- Создание таблицы ---
    table = doc.add_table(rows=2, cols=9)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row in table.rows:
        row.cells[0].width = Inches(2.5)

    set_repeat_table_header(table.rows[0])
    set_repeat_table_header(table.rows[1])

    cell = table.cell(0, 0)
    cell.text = "Товары"
    merge_cells_vertically(table, 0, 0, 1)
    set_cell_background(cell)
    center_cell(cell)
    make_bold(cell)

    cell = table.cell(0, 1)
    cell.text = f"{month_ranges[month]}\n{year-1} года"
    merge_cells_horizontally(table, 0, 1, 3)
    for i in range(1, 4):
        set_cell_background(table.cell(0, i))
        center_cell(table.cell(0, i))
        make_bold(table.cell(0, i))

    cell = table.cell(0, 4)
    cell.text = f"{month_ranges[month]}\n{year} года"
    merge_cells_horizontally(table, 0, 4, 6)
    for i in range(4, 7):
        set_cell_background(table.cell(0, i))
        center_cell(table.cell(0, i))
        make_bold(table.cell(0, i))

    cell = table.cell(0, 7)
    cell.text = f"Прирост\n{year}/{year-1}"
    merge_cells_horizontally(table, 0, 7, 8)
    for i in range(7, 9):
        set_cell_background(table.cell(0, i))
        center_cell(table.cell(0, i))
        make_bold(table.cell(0, i))

    headers = ["физ. объем", "тыс.$", "Доля", "физ. объем", "тыс.$", "Доля", "физ. объем", "млн.$"]
    for i in range(8):
        cell = table.cell(1, i + 1)
        cell.text = headers[i]
        set_cell_background(cell)
        center_cell(cell)
        make_bold(cell)

    # --- Добавление данных ---
    for i, row_data in enumerate(table_data):
        row = table.add_row().cells
        for col, val in enumerate(row_data):
            cell = row[col]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.runs[0]
                run.font.name = 'Times New Roman'
                run.font.size = Pt(9)

                if i == 0:
                    run.bold = True
                else:
                    if col in [1, 4]:  # 2 и 5 столбец — жирный
                        run.bold = True
                    if col in [2, 5]:  # 3 и 6 столбец — курсив
                        run.italic = True
                    if col in [7, 8]:  # 7 и 8 столбец — цвет по значению
                        if val.strip().lower() == "new":
                            run.font.color.rgb = RGBColor(0, 176, 80)  # зелёный
                        elif val.strip().startswith("-"):
                            run.font.color.rgb = RGBColor(255, 0, 0)  # красный



# Финальная сборка документа
def generate_report(filename='final_report3.docx'):
    doc = Document()

    # Установка отступов страницы (в дюймах)
    section = doc.sections[0]
    section.top_margin = Inches(2 / 2.54)
    section.bottom_margin = Inches(2 / 2.54)
    section.left_margin = Inches(3 / 2.54)
    section.right_margin = Inches(1.5 / 2.54)

    add_document_header(doc, data_for_doc["document_header"])
    add_summary_paragraph(doc, data_for_doc["summary_text"])
    add_summary_table(doc, data_for_doc["summary_header"], data_for_doc["summary_table"])
    
    add_import_analysis_text(doc, data_for_doc["import_text"])
    add_import_analysis_text(doc, data_for_doc["export_text"])

    generate_export_doc_with_repeating_header(doc, data_for_doc["export_header"], data_for_doc["export_table"][:15], month, year, units_of_account)
    generate_export_doc_with_repeating_header(doc, data_for_doc["import_header"], data_for_doc["import_table"][:15], month, year, units_of_account)

    doc.save(filename)

# Генерация документа
generate_report()
